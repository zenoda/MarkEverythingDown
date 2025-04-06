from processors.base import BaseDocumentProcessor, StructuredDocument, DocumentSection, DocumentElement, DocumentType
from pathlib import Path
import mammoth
import re
import docx
from docx2python import docx2python
import pandas as pd
from io import StringIO

class DocxProcessor(BaseDocumentProcessor):
    """Processor for DOCX documents with enhanced table support"""
    
    def process(self, file_path: str) -> StructuredDocument:
        """Process DOCX document with improved table handling"""
        document = StructuredDocument(
            title=Path(file_path).stem,
            source_file=file_path,
            doc_type=DocumentType.WORD
        )
        
        try:
            # First approach: Try to extract tables using python-docx directly
            doc = docx.Document(file_path)
            has_tables = len(doc.tables) > 0
            
            # Convert DOCX to Markdown using mammoth for general content
            with open(file_path, "rb") as docx_file:
                result = mammoth.convert_to_markdown(docx_file)
            markdown_content = result.value
            
            # If document has tables, use our direct table extraction
            if (has_tables):
                markdown_content = self._handle_tables_directly(doc, markdown_content)
            else:
                # As a fallback, try docx2python for table extraction
                try:
                    doc_data = docx2python(file_path)
                    if hasattr(doc_data, 'tables') and doc_data.tables:
                        markdown_content = self._process_tables(file_path, markdown_content, doc_data.tables)
                except Exception as table_error:
                    print(f"Warning: Unable to process tables with docx2python: {table_error}")
            
            # Create a single section with the markdown content
            section = DocumentSection(title="Document Content")
            section.add_element(DocumentElement(
                content=markdown_content,
                element_type="markdown"
            ))
            document.add_section(section)
            
            # Store raw markdown in metadata
            document.metadata["markdown"] = markdown_content
            
            # Extract any messages/warnings
            if result.messages:
                document.metadata["conversion_messages"] = [msg.message for msg in result.messages]
                
        except Exception as e:
            print(f"Error extracting content from DOCX: {e}")
            # Create error section
            error_section = DocumentSection(title="Error")
            error_section.add_element(DocumentElement(
                content=f"Failed to process document: {str(e)}",
                element_type="paragraph"
            ))
            document.add_section(error_section)
        
        return document
        
    def _handle_tables_directly(self, doc, content):
        """
        Extract tables directly from python-docx and insert them as proper markdown tables
        
        Args:
            doc: python-docx Document object
            content: Markdown content from mammoth
            
        Returns:
            str: Markdown content with properly formatted tables
        """
        if not doc.tables:
            return content
            
        # Split content into paragraphs
        paragraphs = content.split('\n\n')
        result_paragraphs = []
        table_markers = []
        
        # Find potential table markers
        for i, para in enumerate(paragraphs):
            if 'table' in para.lower():
                table_markers.append(i)
                
        # If no explicit markers, look for potential places by identifying short consecutive lines
        if not table_markers:
            for i in range(len(paragraphs) - 1):
                lines = paragraphs[i].split('\n')
                # Look for potential table headers (short lines)
                if all(len(line) < 30 for line in lines) and len(lines) > 1:
                    table_markers.append(i)
        
        # Process each paragraph, inserting tables at appropriate places
        i = 0
        tables_inserted = 0
        
        while i < len(paragraphs):
            if i in table_markers and tables_inserted < len(doc.tables):
                # Add the current paragraph (table caption or header)
                result_paragraphs.append(paragraphs[i])
                
                # Convert the table to markdown and add it
                table_md = self._table_to_markdown(doc.tables[tables_inserted])
                result_paragraphs.append(table_md)
                
                # Skip any paragraphs that might be part of the table in the original content
                next_text = False
                j = i + 1
                while j < len(paragraphs) and not next_text:
                    # Check if paragraph j looks like normal text
                    if len(paragraphs[j].split('\n')) <= 1 and len(paragraphs[j]) > 30:
                        next_text = True
                    j += 1
                
                if next_text:
                    i = j - 1  # Go back to the text paragraph
                else:
                    i = j
                
                tables_inserted += 1
            else:
                # Regular paragraph, just add it
                result_paragraphs.append(paragraphs[i])
                i += 1
        
        # If we still have tables left, add them at the end
        for j in range(tables_inserted, len(doc.tables)):
            table_md = self._table_to_markdown(doc.tables[j])
            result_paragraphs.append("Additional table found:")
            result_paragraphs.append(table_md)
            
        return '\n\n'.join(result_paragraphs)
        
    def _table_to_markdown(self, table):
        """
        Convert python-docx table to markdown table format
        
        Args:
            table: python-docx Table object
            
        Returns:
            str: Markdown table
        """
        if not table.rows:
            return ""
            
        markdown_rows = []
        
        # Get cell text for each row (preserve empty cells)
        all_rows = []
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                # Remove newlines and extra spaces from cell text
                cell_text = cell.text.strip().replace('\n', ' ')
                cell_text = re.sub(r'\s+', ' ', cell_text)
                row_cells.append(cell_text if cell_text else " ")
            all_rows.append(row_cells)
            
        # Determine the number of columns based on the row with the most cells
        num_cols = max(len(row) for row in all_rows)
        
        # Ensure all rows have the same number of columns
        for row in all_rows:
            while len(row) < num_cols:
                row.append(" ")
        
        # Generate header row
        header_row = "| " + " | ".join(all_rows[0]) + " |"
        markdown_rows.append(header_row)
        
        # Generate separator row
        separator_row = "| " + " | ".join(["---"] * num_cols) + " |"
        markdown_rows.append(separator_row)
        
        # Generate data rows
        for row in all_rows[1:]:
            data_row = "| " + " | ".join(row) + " |"
            markdown_rows.append(data_row)
            
        return "\n".join(markdown_rows)
    
    def _process_tables(self, file_path: str, content: str, tables_data) -> str:
        """
        Process tables extracted by docx2python and replace mammoth's table output
        with better markdown tables
        
        Args:
            file_path: Path to the DOCX file
            content: Markdown content from mammoth
            tables_data: Tables extracted by docx2python
            
        Returns:
            str: Markdown content with improved tables
        """
        if not tables_data:
            # Direct table detection in the content if docx2python didn't find tables
            return self._detect_and_format_tables_from_content(content, file_path)
        
        try:
            # Try to use python-docx to get additional table information
            doc = docx.Document(file_path)
            
            # Find table markers in mammoth output
            table_pattern = r'\|[^\n]+\|\n\|[^\n]+\|'
            table_matches = list(re.finditer(table_pattern, content))
            
            # If no tables found with the standard pattern, try a broader pattern
            if not table_matches:
                # This broader pattern looks for possible table content by finding consecutive lines 
                # with keyword patterns that might indicate a flattened table
                broader_pattern = r'(?:[\w\s]+\n){2,}(?:[\w\s]+)'
                table_matches = list(re.finditer(broader_pattern, content))
            
            # If we still don't find tables with regex patterns, use doc.tables directly
            if (not table_matches or not doc.tables) and doc.tables:
                return self._reconstruct_tables_from_python_docx(content, doc.tables)
                
            # Start with the original content
            result = content
            
            # If tables found in docx but not in mammoth output, reconstruct them completely
            if doc.tables and not table_matches:
                return self._reconstruct_tables_from_python_docx(content, doc.tables)
                
            # Standard processing when we have matching tables
            offset = 0
            
            # Process each table
            for i, table_data in enumerate(tables_data):
                if i >= len(doc.tables):
                    break
                    
                # Get table from python-docx
                table = doc.tables[i]
                
                # Find the table in the content
                if i < len(table_matches):
                    table_match = table_matches[i]
                    start_pos = table_match.start() + offset
                    end_pos = table_match.end() + offset
                else:
                    # If we don't have an exact match, place the table at the end
                    start_pos = len(result)
                    end_pos = start_pos
                
                # Convert table to pandas DataFrame
                df = self._table_to_dataframe(table_data)
                
                # Convert DataFrame to markdown table
                markdown_table = self._df_to_markdown(df)
                
                # Replace table in content
                result = result[:start_pos] + markdown_table + result[end_pos:]
                
                # Update offset for next replacement
                offset += len(markdown_table) - (end_pos - start_pos)
            
            return result
        
        except Exception as e:
            print(f"Error processing tables: {e}")
            # Fallback to direct content processing
            return self._detect_and_format_tables_from_content(content, file_path)
    
    def _detect_and_format_tables_from_content(self, content: str, file_path: str) -> str:
        """
        Detect and format tables directly from the content when docx2python fails
        
        Args:
            content: Markdown content
            file_path: Path to the DOCX file
            
        Returns:
            str: Content with improved tables
        """
        try:
            # Try to open with python-docx directly
            doc = docx.Document(file_path)
            if not doc.tables:
                return content
                
            return self._reconstruct_tables_from_python_docx(content, doc.tables)
        except:
            # Another fallback approach - try to detect tables based on text patterns
            # Look for lines that might be table headers (consecutive short lines followed by a series of values)
            lines = content.split('\n')
            result = []
            i = 0
            while i < len(lines):
                line = lines[i]
                # Check if this line might be the start of a table (contains "table" keyword)
                if "table" in line.lower() and i + 2 < len(lines):
                    potential_headers = []
                    j = i + 1
                    # Collect potential header lines (short lines with no punctuation)
                    while j < len(lines) and len(lines[j].strip()) < 30 and not re.search(r'[.,:;]', lines[j]) and lines[j].strip():
                        potential_headers.append(lines[j].strip())
                        j += 1
                    
                    # If we found potential headers, format as a table
                    if len(potential_headers) >= 2:
                        result.append(line)  # Add the table description line
                        result.append("")    # Empty line before table
                        
                        # Create markdown table headers
                        result.append("| " + " | ".join(potential_headers) + " |")
                        result.append("| " + " | ".join(["---"] * len(potential_headers)) + " |")
                        
                        # Try to detect table rows
                        k = j
                        row_data = []
                        while k < len(lines) and len(row_data) < 20:  # Limit to 20 rows to avoid false positives
                            if not lines[k].strip():
                                k += 1
                                continue
                            
                            if len(lines[k].strip()) < 30:
                                row_data.append(lines[k].strip())
                            else:
                                break  # Stop if we hit a long line (paragraph)
                            k += 1
                        
                        # Create rows based on headers (split data into rows)
                        rows = []
                        row = []
                        for item in row_data:
                            row.append(item)
                            if len(row) == len(potential_headers):
                                rows.append("| " + " | ".join(row) + " |")
                                row = []
                        
                        # If we have a partial row at the end, add it with empty cells
                        if row:
                            while len(row) < len(potential_headers):
                                row.append("")
                            rows.append("| " + " | ".join(row) + " |")
                        
                        # Add all rows to result
                        result.extend(rows)
                        result.append("")  # Empty line after table
                        
                        i = k  # Skip processed lines
                    else:
                        result.append(line)
                        i += 1
                else:
                    result.append(line)
                    i += 1
            
            return "\n".join(result)
    
    def _reconstruct_tables_from_python_docx(self, content: str, tables) -> str:
        """
        Completely reconstruct tables using python-docx
        
        Args:
            content: Original markdown content
            tables: Tables from python-docx
            
        Returns:
            str: Content with properly formatted tables
        """
        # First, find places in the content where tables might go
        # Look for lines mentioning "table" for table captions/descriptions
        table_markers = []
        lines = content.split('\n')
        
        for i, line in enumerate(lines):
            if re.search(r'table|Table', line):
                table_markers.append(i)
        
        # If we found markers for all tables, use them to insert tables
        if table_markers and len(table_markers) >= len(tables):
            for i, table in enumerate(tables):
                if i >= len(table_markers):
                    break
                    
                marker_pos = table_markers[i]
                
                # Create markdown table from this table
                markdown_table = self._convert_python_docx_table(table)
                
                # Insert after the marker line
                lines.insert(marker_pos + 1, "")
                lines.insert(marker_pos + 2, markdown_table)
                lines.insert(marker_pos + 3, "")
                
                # Update remaining marker positions
                for j in range(i + 1, len(table_markers)):
                    table_markers[j] += 3
            
            return "\n".join(lines)
        else:
            # If we can't find markers for all tables, append tables at the end
            result = content
            
            for table in tables:
                markdown_table = self._convert_python_docx_table(table)
                result += "\n\n" + markdown_table + "\n"
                
            return result
    
    def _convert_python_docx_table(self, table) -> str:
        """
        Convert a python-docx table to markdown
        
        Args:
            table: python-docx table object
            
        Returns:
            str: Markdown table
        """
        # Extract the table data
        table_data = []
        
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                # Get text and clean it
                text = cell.text.strip().replace('\n', ' ')
                row_data.append(text)
            table_data.append(row_data)
            
        # If no data, return empty string
        if not table_data:
            return ""
            
        # Get the max columns
        max_cols = max(len(row) for row in table_data)
        
        # Ensure all rows have the same number of columns
        for row in table_data:
            while len(row) < max_cols:
                row.append("")
        
        # Create header row
        header = "| " + " | ".join(table_data[0]) + " |"
        
        # Create separator row
        separator = "| " + " | ".join(["---"] * max_cols) + " |"
        
        # Create data rows
        rows = []
        for row in table_data[1:]:
            rows.append("| " + " | ".join(row) + " |")
            
        # Combine all parts
        return header + "\n" + separator + "\n" + "\n".join(rows)
    
    def _df_to_markdown(self, df: pd.DataFrame) -> str:
        """Convert pandas DataFrame to markdown table"""
        if df.empty:
            return ''
            
        # Use pandas built-in to_markdown when possible
        try:
            return df.to_markdown(index=False)
        except Exception as e:
            # Fallback to manual conversion
            return self._complex_df_to_markdown(df)